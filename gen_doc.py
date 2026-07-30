# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right"):
        border = OxmlElement("w:" + edge)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading_cn(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h

def add_para(text, bold=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_table_cn(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)
        set_cell_border(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Microsoft YaHei"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9.5)
            if i == 0:
                run.bold = True
            set_cell_border(cells[i])
    return table

title = doc.add_heading("ToDoWell 2.0.0 技术规格与实现文档", level=0)
for run in title.runs:
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

add_para("版本: 2.0.0   日期: 2026-07-25   作者: 天才的5014")
add_para("本文档记录 ToDoWell 2.0.0 的界面设计、功能、实现方法、编码方式与依赖关系,供后续维护与回溯使用。")

add_heading_cn("一、软件概述", 1)
add_para("ToDoWell 是一款桌面端极简待办事项管理工具,由 Python 3.8 + tkinter 的 1.5.0 版本完全用 C++ 和 Direct2D 重写。启动后常驻桌面右下角,支持项目分类管理、待办任务增删改、开机自启、JSON 持久化、多种归位/关闭动画。")
add_para("目标运行平台: Windows 7 (64 位, 需 SP1 + Platform Update) 与 Windows 11。开发平台为 Windows 11 + Visual Studio 2022, 但实际使用在 Win7。", bold=True)

add_heading_cn("二、窗口与界面规格", 1)
add_heading_cn("2.1 窗口属性", 2)
add_bullet("无系统标题栏的自绘窗口 (WS_POPUP + WS_THICKFRAME),由程序自行绘制标题栏、内容区、底部栏。")
add_bullet("圆角半径 8px。Win11 优先 DwmSetWindowAttribute 系统级圆角;Win7 降级 CreateRoundRectRgn + SetWindowRgn。两种方案先后尝试。")
add_bullet("高 DPI 适配: 入口处 SetProcessDPIAware(),所有布局常量以 DIP(1 DIP = 1/96 英寸)定义,渲染时乘以 dpiScale 转物理像素。")
add_bullet("开机自启通过 HKCU\\Software\\Microsoft\\Windows\\CurrentVersion\\Run,键名 LvZhiFangTodo,无需管理员权限。")

add_heading_cn("2.2 启动尺寸(黄金分割)", 2)
add_para("第一次启动时窗口宽高遵循黄金分割比例,之后用户可自由拉伸:")
add_code("width  = 248 DIP          (物理像素 = 248 * dpiScale)")
add_code("height = 256 * phi DIP    phi = 1.6180339887  -> 约 414 DIP")
add_para("启动后自动定位到桌面工作区右下角(SPI_GETWORKAREA 扣除任务栏)。")
add_para("最小尺寸: 240 x 320 DIP。八方向缩放,5px 热区。标题栏拖拽移动,15px 边缘吸附。")

add_heading_cn("2.3 三段式布局", 2)
add_table_cn(["区域","高度(DIP)","背景色","说明"],[
    ["标题栏","38","#FFEAA7 淡黄","左侧前缀+gogogo!!!, 右侧归位箭头和关闭按钮, 可拖拽"],
    ["内容区","自适应","#f5f5f7 浅灰白","Canvas 滚动, 项目卡片+待办列表"],
    ["底部栏","38","#f5f5f7","左侧 + 新建项目, 右侧 齿轮 设置, 无分割线"],
])

add_heading_cn("2.4 配色方案", 2)
add_table_cn(["常量","色值","用途"],[
    ["BG","#ffffff","项目卡片背景,白色浮层"],
    ["PAGE","#f5f5f7","内容区与底部栏背景"],
    ["TEXT","#1d1d1f","正文文字"],
    ["MUTED","#86868b","占位提示文字"],
    ["ACCENT","#0066cc","唯一强调色, 聚焦边框/悬停/光标"],
    ["SEP","#e5e5e5","分割线"],
    ["CIRCLE","#c6c6c8","待办前空心圆圈"],
    ["TITLE_BG","#FFEAA7","标题栏淡黄"],
    ["TITLE_FG","#2c2c2e","标题栏文字"],
    ["INPUT_BD","#d1d1d6","输入框边框"],
    ["DEL_BTN","#c6c6c8","关闭/删除按钮静态色"],
    ["DEL_HOVER","#ff3b30","删除悬停红色"],
    ["PROJ_COLORS","8 色","项目编号循环: 红#e74c3c 蓝#3498db 绿#2ecc71 橙#f39c12 紫#9b59b6 青#1abc9c 深橙#e67e22 深灰蓝#34495e"],
])

add_heading_cn("2.5 字体规范", 2)
add_para("中文统一 Microsoft YaHei(微软雅黑),英文/UI 符号用 Segoe UI Symbol。")
add_table_cn(["位置","字体","字号/字重","颜色"],[
    ["标题栏 gogogo!!! + 前缀","YaHei","11pt Bold","TITLE_FG #2c2c2e"],
    ["标题栏按钮 -> x","Segoe UI Symbol","11pt","TEXT / TITLE_FG"],
    ["项目名称","YaHei","12pt Bold","TEXT"],
    ["项目编号","Segoe UI Symbol","9pt Bold","白色(背景为项目色)"],
    ["待办任务文字","YaHei","10pt","TEXT"],
    ["待办圆圈","Segoe UI Symbol","16pt","CIRCLE"],
    ["底部 + 和 齿轮","Segoe UI Symbol","15pt","TEXT"],
    ["设置弹窗文字","YaHei","10pt","#333333 / #999999"],
    ["关于标题","YaHei","14pt Bold","#333333"],
    ["版本/版权","YaHei","8pt","#aaaaaa"],
])

add_heading_cn("三、功能清单", 1)
add_para("1. 新建项目: 点击底部 +,出现蓝色左边线输入卡片,回车确认。光标自动跳入新建待办输入。")
add_para("2. 编辑项目名: 双击项目名,进入编辑态,回车保存,ESC 取消。")
add_para("3. 删除项目: 点击项目头部右侧 x,悬停变红,点击直接删除(无二次确认)。")
add_para("4. 新建待办: 点击项目底部新建行,回车追加,光标保持可连续输入。")
add_para("5. 完成待办: 点击前圆圈,悬停仅高亮当前项,点击后带淡出动画消失。")
add_para("6. 编辑待办: 点击任务文字进入编辑,回车保存,ESC 取消。")
add_para("7. 滚动: 鼠标滚轮平滑滚动内容区,带动画缓动。")
add_para("8. 窗口移动: 拖拽标题栏,15px 边缘吸附。")
add_para("9. 窗口缩放: 八方向 5px 热区,最小 240x320。")
add_para("10. 归位: 点击 -> 箭头,5 种动画可选(设置中切换)。")
add_para("11. 关闭: 点击 x,带高级关闭动画淡出退出。")
add_para("12. 开机自启: 设置中切换,即时写入 HKCU 注册表。")
add_para("13. 标题前缀: 设置中编辑最多 5 个汉字,回车保存,默认值'你好'。")
add_para("14. 数据持久化: todos.json + config.json,每次操作即时写入。")

add_heading_cn("3.1 归位动画(snap_anim)", 2)
add_table_cn(["序号","名称","效果","时长"],[
    ["0","无动画","瞬移归位","0.01s"],
    ["1","动画1","当前页面从顶部划出,从右下角底部划入出现(默认)","0.80s"],
    ["2","动画2","归位路径中页面消失,从右下角逐步出现","0.80s"],
    ["3","动画3","弹跳归位","0.70s"],
    ["4","动画4","页面从桌面顶端划出,从右下角划入","0.45s"],
])

add_heading_cn("3.2 关闭动画", 2)
add_para("关闭时窗口先轻微弹起(1.0->1.03),再加速收缩淡出(1.03->0.83),伴随向上漂移和极小倾斜,通过 SetLayeredWindowAttributes 控制整体透明度。使用 1ms 多媒体定时器保证流畅。")

add_heading_cn("四、实现方法", 1)
add_heading_cn("4.1 渲染架构", 2)
add_para("采用 Direct2D 硬件加速渲染。Gfx 类封装 ID2D1HwndRenderTarget,直接绑定主窗口 HWND。每帧 BeginDraw -> 绘制全部 UI -> EndDraw,若返回 D2DERR_RECREATE_TARGET 则重建。所有文字、图形、动画均由 Direct2D 直接绘制,不使用任何标准控件(除一个隐藏 EDIT)。")
add_para("动画驱动: timeSetEvent(1ms) 多媒体定时器 -> PostMessage(WM_TIMER) -> tick(dt) -> InvalidateRect -> WM_PAINT -> render。dt 由 QueryPerformanceCounter 高精度计时。在 Win7 RTX A4000 上可跑满刷新率,动画非常流畅。")

add_heading_cn("4.2 自绘输入与中文输入法(IME)", 2)
add_para("所有文本输入框由 Direct2D 自绘(光标、文字、拼音 composition string)。为保证中文输入法候选框跟随光标,采用'隐藏 EDIT 代理'架构:")
add_bullet("创建一个 1x1 像素、不可见、不绘制(WM_PAINT/WM_ERASEBKGND 阻止)的 EDIT 子控件作为 IME 输入代理。")
add_bullet("EDIT 获得键盘焦点,接收 IME composition 事件;D2D 绘制全部可见内容。")
add_bullet("positionEdit(): 根据文字宽度 measureText(editText + compositionText) 计算 D2D 光标实际屏幕坐标,把 1x1 EDIT 定位到该点。")
add_bullet("positionIME(): ImmGetContext + ImmSetCompositionWindow(CFS_POINT) + ImmSetCandidateWindow,用屏幕坐标设置候选框位置。")
add_bullet("WM_IME_COMPOSITION 中先更新拼音文本再定位,避免位置滞后一帧。")
add_para("重要: 此方案在 Win7 上工作正常,因为 Win7 输入法走 IMM32 机制并遵守 ImmSetCandidateWindow。Win11 微软拼音是纯 TSF,会忽略 IMM32 定位,候选框可能飘到左上角——这是平台差异,不影响 Win7 实际使用。", bold=True)

add_heading_cn("4.3 圆圈高亮交互", 2)
add_para("鼠标移到圆圈上只高亮当前项(通过 m_hovCircPi/m_hovCircTi 精确匹配)。点击完成任务后设置 m_suppressCircleHover 抑制,直到鼠标真正移动超过 3 像素才恢复,避免误高亮下一条。")

add_heading_cn("4.4 数据持久化", 2)
add_para("JSON 自实现解析/序列化(json.cpp),无第三方依赖。文件位于 exe 同目录:")
add_code("todos.json: [{name, todos:[{text, done}]}, ...]")
add_code("config.json: {title_prefix, auto_start, snap_anim}")
add_para("每次增删改即时写入,IO 异常静默捕获。首次运行文件不存在时返回空列表/默认值(前缀'你好',动画=1)。")

add_heading_cn("五、编码方式", 1)
add_heading_cn("5.1 技术栈", 2)
add_table_cn(["项","说明"],[
    ["语言","C++17 (/EHsc /O2)"],
    ["图形","Direct2D + DirectWrite (COM 接口)"],
    ["平台 SDK","Windows 7 目标 (_WIN32_WINNT=0x0601, WINVER=0x0601, _NTDDI_VERSION=0x06010000)"],
    ["运行时","/MT 静态链接 C 运行时, exe 无外部 CRT 依赖"],
    ["字符集","UNICODE / _UNICODE"],
    ["源编码","UTF-8 (/utf-8)"],
    ["构建","Visual Studio 2022 / MSVC 14.51, build.bat 批处理"],
])

add_heading_cn("5.2 源码文件结构", 2)
add_table_cn(["文件","职责"],[
    ["main.cpp","WinMain, 窗口注册/创建, WndProc 消息分发, DPI, 多媒体定时器, IME 消息转发"],
    ["app.cpp","App 类: UI 构建、渲染、布局(rebuildHits)、事件处理、编辑、动画 tick、设置/关于弹窗"],
    ["app.h","App 类声明, AppC 布局常量, EditMode/HitType 枚举, Hit 结构"],
    ["gfx.cpp/h","Gfx 类: D2D 工厂、渲染目标、DirectWrite 字体、绘制原语(fillRect/drawText/measureTextW)、调色板 C::、FontId 枚举"],
    ["storage.cpp/h","load/save_projects, load/save_config, is/set_auto_start, exe_dir, UTF-8 转换"],
    ["json.cpp/h","JsonValue 自实现 JSON 解析与序列化"],
    ["resource.h / app.rc / icon.ico","资源: 应用图标 IDI_APP"],
    ["app.manifest","asInvoker, supportedOS Win7-11, dpiAware, Common-Controls v6, 版本 2.0.0.0"],
    ["build.bat","编译脚本: rc + cl + link"],
])

add_heading_cn("5.3 布局常量速查(AppC, DIP)", 2)
add_code("TITLE_H=38  BOT_H=38  RESIZE_M=5  CORNER_R=8")
add_code("MIN_W=240  MIN_H=320  RATIO_H_W=1.5")
add_code("CONTENT_PAD=10  CONTENT_TOP=8  CARD_GAP=10")
add_code("CARD_INNER=14  CARD_TOP=10  ROW_H=26")
add_code("CIRCLE_R=7  BADGE_H=18  BADGE_MIN_W=22")
add_code("SNAP_GAP=15  SCROLL_STEP=32")

add_heading_cn("六、依赖", 1)
add_para("运行时依赖(全部 Windows 7/11 系统自带, 无需安装):")
add_table_cn(["DLL","用途"],[
    ["d2d1.dll","Direct2D 渲染"],
    ["DWrite.dll","DirectWrite 文本布局"],
    ["user32.dll","窗口、消息、输入"],
    ["gdi32.dll","GDI 辅助(圆角区域等)"],
    ["advapi32.dll","注册表(开机自启)"],
    ["winmm.dll","多媒体定时器(高帧率动画)"],
    ["imm32.dll","输入法候选框定位"],
    ["kernel32.dll","系统基础"],
])
add_para("链接库: d2d1.lib dwrite.lib user32.lib gdi32.lib advapi32.lib winmm.lib imm32.lib")
add_para("C 运行时: /MT 静态链接, 不依赖 VCRedist/ucrtbase。exe 单文件约 395 KB。", bold=True)
add_para("无任何第三方库。无 Python、无 .NET、无 Electron。")

add_heading_cn("七、构建与分发", 1)
add_heading_cn("7.1 构建", 2)
add_code("> cd D:\\ToDoWell2")
add_code("> build.bat")
add_para("需要 Visual Studio 2022(含 MSVC 与 Windows SDK)。输出 ToDoWell.exe(约 395 KB)。")

add_heading_cn("7.2 分发", 2)
add_bullet("只分发 ToDoWell.exe 这一个文件。")
add_bullet("todos.json 和 config.json 是运行时数据, 不要打包——首次运行自动创建。")
add_bullet("默认值: 代办列表空白; 标题前缀'你好'; 归位动画=动画1; 开机自启=关。")
add_bullet("拷到 Win7 任意目录双击运行即可, 无需联网/管理员/额外依赖。")
add_bullet("前提: Win7 64 位 + SP1 + Platform Update(Direct2D 需要)。")

add_heading_cn("7.3 版本", 2)
add_para("内部版本 2.0.0(manifest assemblyIdentity version=2.0.0.0)。关于页显示'版本 2.0.0'、'版权所有@天才的5014'。")

add_heading_cn("八、已知行为与注意事项", 1)
add_bullet("Win11 上中文输入法候选框可能不跟随(纯 TSF 不理会 IMM32 定位); Win7 上正常。")
add_bullet("关闭时若数据文件被占用, 写入静默失败不崩溃。")
add_bullet("圆角在 Win7 使用 GDI 区域裁剪, Win11 使用 DWM 合成。")
add_bullet("第一次启动尺寸为黄金分割, 用户拉伸后不记忆尺寸(每次启动回到右下角默认尺寸)。")

doc.save(r"D:\ToDoWell2\ToDoWell_2.0.0_doc.docx")
print("OK saved")
